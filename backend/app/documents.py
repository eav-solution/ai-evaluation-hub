import multiprocessing
import platform
import re
import time
from io import BytesIO
from zipfile import BadZipFile, ZipFile


DEFAULT_MAX_EXPANDED_BYTES = 100 * 1024 * 1024
DEFAULT_MAX_PAGES = 2000
DEFAULT_MAX_CHARS = 20_000_000


def original_storage_key(workspace_id: str, document_id: str, format: str) -> str:
    return f"documents/{workspace_id}/{document_id}.{format}"


def text_storage_key(workspace_id: str, document_id: str) -> str:
    return f"documents/{workspace_id}/{document_id}.extracted.txt"


def _decode(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("File must be UTF-8 encoded") from exc


def extract_text(
    data: bytes,
    format: str,
    *,
    max_expanded_bytes: int = DEFAULT_MAX_EXPANDED_BYTES,
    max_pages: int = DEFAULT_MAX_PAGES,
    max_chars: int = DEFAULT_MAX_CHARS,
) -> str:
    if format == "pdf":
        from pypdf import PdfReader

        try:
            reader = PdfReader(BytesIO(data))
            if reader.is_encrypted:
                raise ValueError("Encrypted PDFs are not supported")
            if len(reader.pages) > max_pages:
                raise ValueError("PDF exceeds the page limit")
            parts: list[str] = []
            char_count = 0
            for page in reader.pages:
                part = page.extract_text() or ""
                char_count += len(part) + (2 if parts else 0)
                if char_count > max_chars:
                    raise ValueError("Document exceeds the extracted text limit")
                parts.append(part)
            text = "\n\n".join(parts)
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("Could not read PDF") from exc
    elif format == "docx":
        from docx import Document as DocxDocument

        try:
            with ZipFile(BytesIO(data)) as archive:
                expanded_bytes = sum(item.file_size for item in archive.infolist())
            if expanded_bytes > max_expanded_bytes:
                raise ValueError("DOCX exceeds the expanded size limit")
            document = DocxDocument(BytesIO(data))
        except ValueError:
            raise
        except BadZipFile as exc:
            raise ValueError("Could not read DOCX") from exc
        except Exception as exc:
            raise ValueError("Could not read DOCX") from exc
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n\n".join(parts)
    elif format == "html":
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(_decode(data), "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    elif format in {"txt", "md"}:
        text = _decode(data)
    else:
        raise ValueError(f"Unsupported document format: {format}")

    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text).strip()
    if len(text) > max_chars:
        raise ValueError("Document exceeds the extracted text limit")
    if not text:
        raise ValueError("No extractable text (scanned or empty document?)")
    return text


def _extract_text_process(
    connection,
    data: bytes,
    format: str,
    max_expanded_bytes: int,
    max_pages: int,
    max_chars: int,
    max_memory_bytes: int,
) -> None:
    try:
        if platform.system() == "Linux":
            import resource

            resource.setrlimit(
                resource.RLIMIT_AS, (max_memory_bytes, max_memory_bytes)
            )
        text = extract_text(
            data,
            format,
            max_expanded_bytes=max_expanded_bytes,
            max_pages=max_pages,
            max_chars=max_chars,
        )
        connection.send((True, text))
    except BaseException as exc:
        try:
            connection.send((False, str(exc)))
        except Exception:
            pass
    finally:
        connection.close()


def extract_text_isolated(
    data: bytes,
    format: str,
    *,
    max_expanded_bytes: int = DEFAULT_MAX_EXPANDED_BYTES,
    max_pages: int = DEFAULT_MAX_PAGES,
    max_chars: int = DEFAULT_MAX_CHARS,
    timeout_seconds: float = 30,
    max_memory_bytes: int = 512 * 1024 * 1024,
) -> str:
    context = multiprocessing.get_context("spawn")
    parent, child = context.Pipe(duplex=False)
    process = context.Process(
        target=_extract_text_process,
        args=(
            child,
            data,
            format,
            max_expanded_bytes,
            max_pages,
            max_chars,
            max_memory_bytes,
        ),
    )
    process.start()
    child.close()
    deadline = time.monotonic() + timeout_seconds
    while not parent.poll(0.05):
        if not process.is_alive():
            parent.close()
            raise ValueError("Document parser process failed")
        if time.monotonic() >= deadline:
            process.terminate()
            process.join()
            parent.close()
            raise ValueError("Document parsing timed out")
    succeeded, result = parent.recv()
    parent.close()
    process.join(1)
    if process.is_alive():
        process.terminate()
        process.join()
    if not succeeded:
        raise ValueError(result or "Document parser process failed")
    return result


MIN_CHUNK_CHARS = 50


def _split_long(paragraph: str, chunk_chars: int) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+", paragraph)
    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        while len(sentence) > chunk_chars:
            if current:
                pieces.append(current)
                current = ""
            pieces.append(sentence[:chunk_chars])
            sentence = sentence[chunk_chars:]
        if not sentence:
            continue
        if current and len(current) + len(sentence) + 1 > chunk_chars:
            pieces.append(current)
            current = sentence
        else:
            current = f"{current} {sentence}" if current else sentence
    if current:
        pieces.append(current)
    return pieces


def chunk_text(text: str, chunk_chars: int) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    pieces: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) > chunk_chars:
            pieces.extend(_split_long(paragraph, chunk_chars))
        else:
            pieces.append(paragraph)
    chunks: list[str] = []
    current = ""
    for piece in pieces:
        if current and len(current) + len(piece) + 2 > chunk_chars:
            chunks.append(current)
            current = piece
        else:
            current = f"{current}\n\n{piece}" if current else piece
    if current:
        chunks.append(current)
    return [chunk for chunk in chunks if len(chunk) >= MIN_CHUNK_CHARS]
